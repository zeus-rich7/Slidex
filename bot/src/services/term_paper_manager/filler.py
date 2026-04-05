from dataclasses import dataclass
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches

from src.helpers.async_factory import to_async


# json_response = {'outline': {'introduction': 'Kirish', 'sections': ['1. Bank marketingining mohiyati, maqsad va vazifalari', '2. Bank marketingi strategiyasi va marketing-miks (7P) tushunchasi', '3. Bank bozorini segmentlash va mijozlar bilan ishlash tamoyillari', "4. O'zbekiston bank tizimida raqamli marketing va innovatsiyalar", '5. Bank xizmatlari sifatini oshirish va mijozlar sodiqligini boshqarish'], 'conclusion': 'Xulosa', 'references': 'Foydalanilgan adabiyotlar'}, 'introduction': "Bugungi kunda jahon iqtisodiyotining globallashuvi va moliya bozorlaridagi raqobatning kuchayishi sharoitida bank marketingi tijorat banklari faoliyatining ajralmas qismiga aylandi. O'zbekiston Respublikasida amalga oshirilayotgan bank-moliya tizimidagi islohotlar, xususan, banklarni xususiylashtirish va transformatsiya qilish jarayonlari mijozlarga yo'naltirilgan zamonaviy marketing yondashuvlarini talab etmoqda. Mazkur mavzuning dolzarbligi shundaki, an'anaviy bank xizmatlari o'rnini yuqori texnologik va masofaviy xizmatlar egallayotgan bir paytda, banklarning bozordagi o'rnini saqlab qolishi va daromadlilik darajasini oshirishi bevosita marketing strategiyasining samaradorligiga bog'liqdir.\n\nO'zbekiston bank tizimini rivojlantirishning 2020-2025-yillarga mo'ljallangan strategiyasida bank xizmatlari sifatini oshirish va aholining bank tizimiga bo'lgan ishonchini mustahkamlash ustuvor vazifa etib belgilangan. Bu esa, o'z navbatida, bank marketingi vositalaridan foydalangan holda mijozlar ehtiyojlarini chuqur o'rganish, yangi turdagi innovatsion mahsulotlarni yaratish va ularni bozorda samarali ilgari surishni taqozo etadi. Referatda bank marketingining nazariy asoslari, uning o'ziga xos xususiyatlari hamda O'zbekiston sharoitida qo'llanilish istiqbollari atroflicha tahlil qilinadi.", 'sections': [{'title': '1. Bank marketingining mohiyati, maqsad va vazifalari', 'content': "Bank marketingi — bu bankning mavjud va potensial mijozlari ehtiyojlarini aniqlash, qondirish va ularni ushlab turish orqali foyda olishga yo'naltirilgan boshqaruv jarayonidir. Umumiy marketingdan farqli o'laroq, bank marketingi moliyaviy xizmatlarning nomoddiy xarakterga egaligi, vaqt va ishonch omillariga bog'liqligi bilan ajralib turadi. Bank marketingining asosiy maqsadi — bank resurslarini mijozlar talabiga muvofiq samarali joylashtirish, bozordagi ulushni kengaytirish va uzoq muddatli barqarorlikni ta'minlashdan iborat.\n\nBank marketingining vazifalari kompleks xarakterga ega bo'lib, ular qatoriga bozor kon'yunkturasini tahlil qilish, raqobatchilar faoliyatini o'rganish, yangi bank mahsulotlarini ishlab chiqish va ularning narx siyosatini belgilash kiradi. Shuningdek, bankning ijobiy imijini shakllantirish va har bir mijoz bilan individual ishlash tizimini yo'lga qo'yish muhim vazifalardan sanaladi. Hozirgi davrda bank marketingi shunchaki xizmatlarni sotish emas, balki mijozlar uchun qo'shimcha qiymat yaratish falsafasiga asoslanadi. Bu jarayonda bank nafaqat moliyaviy vositachi, balki mijozning biznes hamkoriga aylanadi. O'zbekistonning tijorat banklari, masalan, 'TBC Bank' yoki 'Anorbank' kabi raqamli banklar misolida ko'rishimiz mumkinki, marketing yondashuvi butunlay mijoz tajribasiga (Customer Experience) yo'naltirilgan. Bu esa an'anaviy banklarni ham o'z ish uslublarini qayta ko'rib chiqishga majbur qilmoqda. Marketing tadqiqotlari orqali banklar qaysi segmentlarda yuqori daromad mavjudligini va qaysi hududlarda xizmatlarga talab ortib borayotganini aniqlaydilar, bu esa risklarni boshqarishda ham muhim rol o'ynaydi."}, {'title': '2. Bank marketingi strategiyasi va marketing-miks (7P) tushunchasi', 'content': "Bank marketingi strategiyasi bankning umumiy biznes rejasining tarkibiy qismi bo'lib, u uzoq muddatli rivojlanish yo'nalishlarini belgilaydi. Bank xizmatlari bozorida muvaffaqiyatga erishish uchun marketing-miks elementlarini to'g'ri uyg'unlashtirish zarur. Agar an'anaviy ishlab chiqarishda 4P modeli qo'llanilsa, xizmat ko'rsatish sohasida, jumladan banklarda 7P modeli (Product, Price, Place, Promotion, People, Process, Physical Evidence) samaraliroq hisoblanadi.\n\n1. Product (Mahsulot) — bank taqdim etayotgan barcha xizmatlar (kreditlar, depozitlar, kartalar). 2. Price (Narx) — foiz stavkalari, komissiyalar va xizmat haqlari. 3. Place (Joy) — bank filiallari tarmog'i va masofaviy xizmat ko'rsatish kanallari. 4. Promotion (Rag'batlantirish) — reklama, PR va sotuvni rag'batlantirish tadbirlari. 5. People (Xodimlar) — bank xodimlarining malakasi va mijozlar bilan muloqot madaniyati. 6. Process (Jarayon) — xizmat ko'rsatish tezligi va tartibi. 7. Physical Evidence (Moddiy dalillar) — bank binosining ko'rinishi, logotipi va brendning tan olinishi. O'zbekiston banklarida narx omili uzoq vaqt asosiy raqobat quroli bo'lib kelgan bo'lsa, hozirda 'Process' va 'People' omillariga e'tibor kuchaymoqda. Mijozlar endilikda nafaqat past foizli kredit, balki tezkor xizmat ko'rsatish (skoring tizimlari) va malakali maslahat kutmoqdalar. Marketing-miksning barcha elementlari o'zaro muvofiq bo'lgandagina bank barqaror raqobatbardoshlikka erisha oladi. Strategiyani ishlab chiqishda SWOT-tahlil usulidan foydalanib, bankning kuchli va kuchsiz tomonlari, tashqi imkoniyatlar va tahdidlar baholanadi. Bu esa o'zgaruvchan iqtisodiy sharoitda egiluvchan marketing siyosatini yuritish imkonini beradi."}, {'title': '3. Bank bozorini segmentlash va mijozlar bilan ishlash tamoyillari', 'content': "Bozorni segmentlash bank marketingining fundamental asoslaridan biri bo'lib, u bank xizmatlariga bo'lgan ehtiyojlari o'xshash bo'lgan mijozlar guruhlarini ajratishni anglatadi. Banklar barcha mijozlarga bir xil xizmat ko'rsata olmaydi, shuning uchun maqsadli segmentlarni aniqlash resurslarni tejash va samaradorlikni oshirishga xizmat qiladi. Odatda segmentlash geografik, demografik, iqtisodiy va xulq-atvor mezonlari asosida amalga oshiriladi.\n\nKorporativ mijozlar segmentida banklar yirik korxonalar, kichik va o'rta biznes subyektlariga bo'linadi. Chakana segmentda esa daromad darajasi, yoshi va hayot tarzi asosiy mezon bo'ladi. Masalan, O'zbekistonda yoshlar uchun maxsus 'talaba kartalari' yoki tadbirkor ayollar uchun imtiyozli kredit liniyalari aynan segmentlash natijasidir. Maqsadli segment tanlangandan so'ng, bank ushbu segmentda o'z 'pozitsiyasini' egallashi kerak. Pozitsiyalash — bu mijoz ongida bank haqida ijobiy va o'ziga xos tasavvurni shakllantirishdir. Hozirgi kunda 'mijozga yo'naltirilganlik' (customer-centricity) tamoyili banklar uchun ustuvor ahamiyat kasb etmoqda. Bu har bir mijozning ehtiyojlarini tushunish va unga moslashtirilgan (customized) mahsulotlarni taklif qilishni nazarda tutadi. Mijozlar bilan ishlashda CRM (Customer Relationship Management) tizimlari muhim rol o'ynaydi. Ushbu tizim orqali bank mijozning barcha tranzaksiyalari, qiziqishlari va shikoyatlarini tahlil qilib, unga kerakli vaqtda kerakli mahsulotni taklif qila oladi. Masalan, mijoz doimiy ravishda onlayn xaridlar qilsa, bank unga keshbek funksiyasi yuqori bo'lgan kartani taklif etishi maqsadga muvofiqdir. Bunday yondashuv mijozning bankka bo'lgan sadoqatini oshiradi va o'zaro manfaatli hamkorlikni mustahkamlaydi."}, {'title': "4. O'zbekiston bank tizimida raqamli marketing va innovatsiyalar", 'content': "Global trendlarga mos ravishda, O'zbekiston bank sektorida raqamli marketing jadal rivojlanmoqda. An'anaviy reklama vositalari (televideniye, gazeta, bannerlar) o'rnini ijtimoiy tarmoqlar, qidiruv tizimlari va mobil ilovalar egallamoqda. Raqamli marketing banklarga minimal xarajatlar bilan keng auditoriyani qamrab olish va natijalarni aniq o'lchash imkonini beradi. Bugungi kunda 'Ipak Yo'li Banki', 'Hamkorbank' va 'Kapitalbank' kabi muassasalar SMM va kontekstli reklamadan faol foydalanmoqdalar.\n\nInnovatsiyalar bank marketingining harakatlantiruvchi kuchi hisoblanadi. Bunga mobil banking, internet banking, chatbotlar va sun'iy intellekt asosidagi xizmatlarni misol keltirish mumkin. O'zbekistonda bank xizmatlarini raqamlashtirish mijozlarga bankka bormasdan turib hisob raqami ochish, kredit rasmiylashtirish va omonat qo'yish imkoniyatini berdi. Marketing nuqtai nazaridan bu 'xizmat ko'rsatish nuqtasi' tushunchasini o'zgartirdi — endi bank mijozning smartfonida joylashgan. Shuningdek, Big Data (katta ma'lumotlar) tahlili banklarga mijozlar xatti-harakatlarini bashorat qilish imkonini bermoqda. Masalan, sun'iy intellekt mijozning kredit to'lash qobiliyatini bir necha soniya ichida tahlil qilib (skoring), unga mos kredit limitini taklif qiladi. Bank marketingida 'Omnichannel' yondashuvi ham muhim bo'lib, u barcha aloqa kanallari (filial, ilova, call-markaz) o'rtasida uzluksizlikni ta'minlaydi. O'zbekiston bozoriga yangi kirib kelgan 'neobanklar' faqat raqamli marketing orqali o'z mijozlar bazasini kengaytirmoqda, bu esa bozordagi raqobat muhitini yanada sog'lomlashtiradi. Raqamli innovatsiyalar nafaqat qulaylik yaratadi, balki banklar uchun operatsion xarajatlarni kamaytirish va xizmatlar tannarxini pasaytirish imkonini ham beradi."}, {'title': '5. Bank xizmatlari sifatini oshirish va mijozlar sodiqligini boshqarish', 'content': "Bank xizmatlarining o'xshashligi sharoitida xizmat ko'rsatish sifati bankning asosiy raqobatbardoshlik omiliga aylanadi. Xizmat sifati bu — mijozning kutganlari va amalda olgan xizmati o'rtasidagi muvofiqlikdir. Bank marketingida xizmat sifatini baholash uchun ko'pincha SERVQUAL modeli qo'llaniladi, u ishonchlilik, sezgirlik, ishonch, hamdardlik va moddiy dalillar kabi mezonlarni qamrab oladi. Sifatni boshqarish bank xodimlarining muntazam treninglaridan boshlanadi, chunki xizmat ko'rsatish jarayonida inson omili hal qiluvchi rol o'ynaydi.\n\nMijozlar sodiqligi (loyalty) bankning uzoq muddatli foydasi garovidir. Tadqiqotlar shuni ko'rsatadiki, yangi mijozni jalb qilish mavjud mijozni saqlab qolishdan 5-10 baravar qimmatga tushadi. Shuning uchun banklar sodiqlik dasturlarini (cashback, bonus tizimlari, kofeyntlar bilan hamkorlik) ishlab chiqadilar. O'zbekiston banklarida ham turli keshbek dasturlari va ballar yig'ish tizimlari ommalashmoqda. Biroq, haqiqiy sodiqlik faqat moddiy rag'bat bilan emas, balki emotsional bog'liqlik va ishonch bilan shakllanadi. Mijozning shikoyatlari bilan ishlash va qayta aloqani (feedback) yo'lga qo'yish marketingning muhim qismidir. Bank mijozning muammosini qanchalik tez va professional tarzda hal qilsa, uning sodiqligi shunchalik ortadi. 'Mijoz ovozi' (Voice of the Customer) dasturlari orqali banklar o'z xizmatlaridagi kamchiliklarni aniqlaydilar va ularni bartaraf etadilar. Xulosa qilib aytganda, sifatli xizmat va mijozlar bilan uzoq muddatli munosabatlarni boshqarish bank marketingining yakuniy natijasi — barqaror brend va yuqori daromadlilikni ta'minlaydi."}], 'conclusion': "Xulosa o'rnida aytish mumkinki, zamonaviy bank marketingi shunchaki reklama emas, balki bankni boshqarishning yaxlit falsafasi va strategik vositasidir. O'zbekiston bank tizimidagi jadal o'zgarishlar, xususiylashtirish jarayonlari va raqamli transformatsiya bank marketingining ahamiyatini yanada oshirmoqda. Tahlillar shuni ko'rsatadiki, muvaffaqiyatli banklar nafaqat keng xizmatlar turini taklif qilmoqda, balki har bir mijoz bilan individual muloqot o'rnatishga, ularning ehtiyojlarini oldindan bashorat qilishga intilmoqda.\n\nO'zbekiston tijorat banklari uchun marketing samaradorligini oshirish bo'yicha quyidagi tavsiyalarni berish mumkin: birinchidan, mijozlar tajribasini (UX/CX) yaxshilash uchun raqamli kanallarni yanada rivojlantirish; ikkinchidan, Big Data tahlilidan foydalangan holda shaxsiylashtirilgan takliflarni ishlab chiqish; uchinchidan, xodimlarning marketing madaniyati va mijozlar bilan ishlash mahoratini doimiy oshirib borish. Kelajakda banklar o'rtasidagi raqobat faqat texnologiyalar darajasida emas, balki mijozlar ishonchini qozonish va ularga haqiqiy qiymat taqdim etish borasida kechadi. Kuchli marketing strategiyasiga ega banklargina bozorda yetakchi o'rinlarni egallaydi va mamlakat iqtisodiy barqarorligiga o'z hissasini qo'shadi.", 'references': ["1. O'zbekiston Respublikasining 'Banklar va bank faoliyati to'g'risida'gi Qonuni. - Toshkent, 2019.", "2. O'zbekiston Respublikasi Prezidentining 2020-yil 12-maydagi PF-5992-sonli '2020-2025-yillarga mo'ljallangan O'zbekiston Respublikasining bank tizimini isloh qilish strategiyasi to'g'risida'gi Farmoni.", "3. Abdullaeva Sh.Z. Bank ishi. - Toshkent: 'Iqtisod-Moliya', 2017. - 450 b.", '4. Kotler P., Keller K.L. Marketing Management. 15th Edition. - Pearson Education, 2016. - 832 p.', '5. Lavrushin O.I. Bankovskiy menedjment. - Moskva: KNORUS, 2016. - 560 s.', '6. Beloglazova G.N., Krolivetskaya L.P. Bankovskoye delo: uchebnik. - Sankt-Peterburg: Piter, 2014. - 400 s.', "7. Khabibullaev I. Bank marketingi: o'quv qo'llanma. - Toshkent: TDIU, 2020. - 210 b.", '8. Meidan A. Quantitative Methods in Banking. - London: Macmillan Press, 1996. - 240 p.', '9. Suzdaltseva L.P. Marketing v bankovskoy sfere. - Moskva: INFRA-M, 2018. - 320 s.', '10. Jallo K., Jallo J. Digital Banking: A Guide to Digital Business Transformation. - Palgrave Macmillan, 2018. - 280 p.', "11. O'zbekiston Respublikasi Markaziy banki rasmiy sayti - www.cbu.uz", '12. The Banker jurnali materiallari - www.thebanker.com']}







@dataclass
class TermPaper:
    university: str
    subject: str
    theme: str
    json_content: dict
    output_path: str
    def __post_init__(self):
        self.doc = Document()
        self._set_margins()

    # TODO: this should be protected
    def add_text(self, text: str, alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT, font_size: int = 14,
                 font_name: str = "Times New Roman",  bold: bool = False, first_line_indent: Optional[float] = None,
                 space_before: int = 0) -> None:
        paragraph = self.doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(space_before)
        paragraph.paragraph_format.first_line_indent = first_line_indent
        paragraph.alignment = alignment
        run = paragraph.add_run(text)
        run.font.size = Pt(font_size)
        run.bold = bold
        run.font.name = font_name

    # TODO: make the function async
    def _set_margins(self):
        for section in self.doc.sections:
            section.left_margin = Inches(1.2)
            section.right_margin = Inches(0.6)
    @to_async()
    def fill(self):
        self.add_text(
            text="O'ZBEKISTON RESPUBLIKASI \nOLIY VA O'RTA MAXSUS TA'LIM VAZIRLIGI",
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            bold=True,
            font_size=20,
        )
        self.add_text(
            text=f"{self.university.upper()} {self.subject.upper()} FANIDAN",
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            space_before=70,
            bold=True,
            font_size=20,
        )
        self.add_text(
            text="REFERATI",
            space_before=96,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            bold=True,
            font_size=48,
        )
        self.add_text(
            text="Tayyorladi: ________________________",
            space_before=108,
            alignment=WD_ALIGN_PARAGRAPH.RIGHT,
            font_size=18,
        )
        self.add_text(
            text="Qabul qildi: ________________________",
            alignment=WD_ALIGN_PARAGRAPH.RIGHT,
            font_size=18,
        )

        self.add_text(
            text="2026",
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            space_before=50,
            bold=True,
            font_size=10,
        )
        self.doc.add_page_break()
        self.add_text(
            text=self.theme.upper(),
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            bold=True,
        )
        self.add_text(
            text="Reja:",
            bold=True,
        )
        self.add_text(
            text="Kirish",
            bold=True,
        )
        for plan in self.json_content["outline"]["sections"]:
            self.add_text(
                text="    " + plan,
            )
        self.add_text(
            text="Xulosa",
            bold=True
        )
        self.add_text(
            text="Foydalanilgan adabiyotlar",
            bold=True
        )
        self.doc.add_page_break()
        self.add_text(
            text="Kirish",
            bold=True,
            alignment=WD_ALIGN_PARAGRAPH.CENTER
        )
        for paragraph in self.json_content["introduction"].splitlines():
            self.add_text(
                text=paragraph,
                first_line_indent=Inches(0.4)
            )
        for section in self.json_content["sections"]:
            self.doc.add_page_break()
            self.add_text(
                text=section["title"],
                bold=True
            )
            self.add_text(
                text=section["content"],
                first_line_indent=Inches(0.4)
            )
        self.doc.add_page_break()
        self.add_text(
            text="Xulosa",
            bold=True
        )
        for paragraph in self.json_content["conclusion"].splitlines():
            self.add_text(
                text=paragraph,
                first_line_indent=Inches(0.4)
            )
        self.doc.add_page_break()
        self.add_text(
            text="Foydalanilgan adabiyotlar",
            bold=True,
            alignment=WD_ALIGN_PARAGRAPH.CENTER
        )
        for reference in self.json_content["references"]:
            self.add_text(
                reference
            )

    @to_async()
    def save(self):
        self.doc.save(self.output_path)


